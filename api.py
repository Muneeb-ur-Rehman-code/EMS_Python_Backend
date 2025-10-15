from fastapi import FastAPI, HTTPException, BackgroundTasks, Request
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import os
import uuid
import logging
import traceback
import tempfile
from docx import Document
from docx2pdf import convert

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI(title="Document Generator API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = r"C:\Users\hp\Downloads\DevRolin_ems_-System-Devrolinems\DevRolin_ems_-System-Devrolinems\Python-Backend\code\templates"
OUTPUT_DIR = os.path.join(BASE_DIR, "generated_docs")
logger.info(f"BASE_DIR: {BASE_DIR}")
logger.info(f"TEMPLATES_DIR: {TEMPLATES_DIR}")
logger.info(f"OUTPUT_DIR: {OUTPUT_DIR}")
os.makedirs(OUTPUT_DIR, exist_ok=True)
logger.info(f"Created OUTPUT_DIR if not exists: {OUTPUT_DIR}")

class OfferLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    DURATION: str
    STARTDATE: str
    SUPNAME: str
    TASKS: str
    POSITION: str
    DEPARTMENT: str
    FROMANDTODATE: str
    TYPE: str
    RESPONSEDATE: str

class TerminationLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    POSITION: str
    TERMDATE: str
    LASTDAY: str

class CertificateData(BaseModel):
    NAME: str
    POSITION: str
    DURATION: str

class AIMLExperienceLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    DURATION: str
    STARTDATE: str
    ENDDATE: str

class WebDevExperienceLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    DURATION: str
    STARTDATE: str
    ENDDATE: str

class GraphicDesignExperienceLetterData(BaseModel):
    REF: str
    DATE: str
    NAME: str
    DURATION: str
    STARTDATE: str
    ENDDATE: str

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    error_detail = f"Error: {str(exc)}\n{traceback.format_exc()}"
    logger.error(f"Global exception caught: {error_detail}")
    return JSONResponse(
        status_code=500,
        content={"detail": str(exc), "trace": traceback.format_exc().split('\n')},
    )


def generate_document(template_name: str, data: dict, output_prefix: str):
    logger.info(f"Starting document generation for template: {template_name}, prefix: {output_prefix}")
    unique_id = str(uuid.uuid4())[:8]
    docx_filename = f"{output_prefix}_{unique_id}.docx"
    pdf_filename = f"{output_prefix}_{unique_id}.pdf"

    docx_path = os.path.join(OUTPUT_DIR, docx_filename)
    pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    logger.info(f"Template path: {template_path}")
    logger.info(f"Output DOCX path: {docx_path}")
    logger.info(f"Output PDF path: {pdf_path}")

    if not os.path.exists(template_path):
        logger.error(f"Template file not found: {template_path}")
        raise HTTPException(status_code=500, detail=f"Template file not found: {template_name}")

    try:
        logger.info(f"Loading template: {template_path}")
        doc = Document(template_path)
        logger.info(f"Template loaded successfully. Processing with data: {data}")

        for p in doc.paragraphs:
            paragraph_text = p.text
            for key in data.keys():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph_text:
                    logger.info(f"Found placeholder {placeholder} in paragraph: {paragraph_text[:50]}...")
            inline = p.runs
            for i in range(len(inline)):
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in inline[i].text:
                        str_value = str(value) if value is not None else ""
                        logger.info(f"Replacing {placeholder} with '{str_value}' in paragraph")
                        inline[i].text = inline[i].text.replace(placeholder, str_value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                str_value = str(value) if value is not None else ""
                                logger.info(f"Replacing {placeholder} with '{str_value}' in table")
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str_value)

        logger.info(f"Saving Word document to {docx_path}")
        doc.save(docx_path)
        logger.info(f"Word document saved successfully: {docx_path}")

        logger.info(f"Converting Word to PDF: {pdf_path}")
        try:
            convert(docx_path, pdf_path)
            logger.info(f"PDF conversion successful: {pdf_path}")
        except Exception as e:
            logger.error(f"PDF conversion failed: {str(e)}\n{traceback.format_exc()}")
            return {
                "docx_path": docx_path,
                "pdf_path": None,
                "docx_filename": docx_filename,
                "pdf_filename": None,
                "warning": "PDF conversion failed, only DOCX is available"
            }

        return {
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "docx_filename": docx_filename,
            "pdf_filename": pdf_filename
        }

    except Exception as e:
        logger.error(f"Document generation failed: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")


def replace_placeholders(doc, replacements):
    logger.info(f"Replacing placeholders in document with: {replacements}")
    for para in doc.paragraphs:
        text = para.text
        for key, value in replacements.items():
            if key in text:
                logger.info(f"Replacing {key} with '{value}' in paragraph")
                text = text.replace(key, value)
        if text != para.text:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    for key, value in replacements.items():
                        if key in text:
                            logger.info(f"Replacing {key} with '{value}' in table cell")
                            text = text.replace(key, value)
                    if text != para.text:
                        for run in para.runs:
                            run.text = ""
                        para.runs[0].text = text
    logger.info("Finished replacing placeholders")

def generate_certificate(template_path, output_pdf, name, position, duration):
    logger.info(f"Generating certificate with template: {template_path}, output: {output_pdf}")
    doc = Document(template_path)
    replacements = {"{{NAME}}": name, "{{POSITION}}": position, "{{DURATION}}": duration}
    logger.info(f"Certificate replacements: {replacements}")
    replace_placeholders(doc, replacements)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_docx = tmp.name
        logger.info(f"Saving temporary DOCX for certificate: {temp_docx}")
        doc.save(temp_docx)

    logger.info(f"Converting certificate DOCX to PDF: {output_pdf}")
    convert(temp_docx, output_pdf)
    logger.info(f"Certificate PDF generated: {output_pdf}")

    os.remove(temp_docx)
    logger.info(f"Temporary DOCX file removed: {temp_docx}")

    return output_pdf

def cleanup_old_files(background_tasks: BackgroundTasks):
    def delete_old_files():
        try:
            logger.info("Starting cleanup of old files")
            files = [os.path.join(OUTPUT_DIR, f) for f in os.listdir(OUTPUT_DIR)]
            files.sort(key=os.path.getctime, reverse=True)
            logger.info(f"Found {len(files)} files in OUTPUT_DIR")
            for old_file in files[40:]:  # Keep 20 documents (both PDF and DOCX)
                if os.path.isfile(old_file):
                    logger.info(f"Removing old file: {old_file}")
                    os.remove(old_file)
            logger.info("Cleanup completed")
        except Exception as e:
            logger.error(f"Cleanup failed: {str(e)}\n{traceback.format_exc()}")
    background_tasks.add_task(delete_old_files)
    logger.info("Scheduled cleanup of old files")

@app.post("/generate/offer-letter")
async def generate_offer_letter(data: OfferLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Received request to generate offer letter for {data.NAME}")
    template_data = {
        "REF": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": str(data.NAME) if data.NAME else "",
        "DURATION": str(data.DURATION) if data.DURATION else "",
        "STARTDATE": str(data.STARTDATE) if data.STARTDATE else "",
        "SUPNAME": str(data.SUPNAME) if data.SUPNAME else "",
        "TASKS": str(data.TASKS) if data.TASKS else "",
        "POSITION": str(data.POSITION) if data.POSITION else "",
        "DEPARTMENT": str(data.DEPARTMENT) if data.DEPARTMENT else "",
        "FROMANDTODATE": str(data.FROMANDTODATE) if data.FROMANDTODATE else "",
        "TYPE": str(data.TYPE) if data.TYPE else "",
        "RESPONSEDATE": str(data.RESPONSEDATE) if data.RESPONSEDATE else ""
    }
    logger.info(f"Template data for offer letter: {template_data}")
    try:
        result = generate_document("offer_template.docx", template_data, "offer_letter")
        logger.info(f"Offer letter generation result: {result}")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Offer letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]
        logger.info(f"Offer letter response: {response}")

        return response

    except Exception as e:
        logger.error(f"Error generating offer letter: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate/termination-letter")
async def generate_termination_letter(data: TerminationLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Received request to generate termination letter for {data.NAME}")
    template_data = {
        "REFNO": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": data.NAME,
        "POSITION": data.POSITION,
        "TERMDATE": data.TERMDATE,
        "LASTDAY": data.LASTDAY
    }
    logger.info(f"Template data for termination letter: {template_data}")
    try:
        result = generate_document("Termination Letter.docx", template_data, "termination_letter")
        logger.info(f"Termination letter generation result: {result}")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Termination letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]
        logger.info(f"Termination letter response: {response}")

        return response
    except Exception as e:
        logger.error(f"Error generating termination letter: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate/certificate")
async def generate_certificate_endpoint(data: CertificateData, background_tasks: BackgroundTasks):
    logger.info(f"Received request to generate certificate for {data.NAME}")
    try:
        unique_id = str(uuid.uuid4())[:8]
        docx_filename = f"certificate_{unique_id}.docx"
        pdf_filename = f"certificate_{unique_id}.pdf"
        docx_path = os.path.join(OUTPUT_DIR, docx_filename)
        pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)
        template_path = os.path.join(TEMPLATES_DIR, "Certificate_Template.docx")
        logger.info(f"Certificate template path: {template_path}")
        logger.info(f"Certificate DOCX path: {docx_path}")
        logger.info(f"Certificate PDF path: {pdf_path}")

        generate_certificate(
            template_path=template_path,
            output_pdf=pdf_path,
            name=data.NAME,
            position=data.POSITION,
            duration=data.DURATION
        )
        logger.info("Certificate generation completed")

        doc = Document(template_path)
        replacements = {
            "{{NAME}}": data.NAME,
            "{{POSITION}}": data.POSITION,
            "{{DURATION}}": data.DURATION
        }
        logger.info(f"Certificate replacements: {replacements}")
        replace_placeholders(doc, replacements)
        logger.info(f"Saving certificate DOCX: {docx_path}")
        doc.save(docx_path)
        logger.info(f"Certificate DOCX saved: {docx_path}")

        cleanup_old_files(background_tasks)
        response = {
            "success": True,
            "message": "Experience certificate generated successfully",
            "docx_url": f"/download/{docx_filename}",
            "pdf_url": f"/download/{pdf_filename}"
        }
        logger.info(f"Certificate response: {response}")

        return response

    except Exception as e:
        logger.error(f"Error generating certificate: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate/experience-letter/aiml")
async def generate_aiml_experience_letter(data: AIMLExperienceLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Received request to generate AI/ML experience letter for {data.NAME}")
    template_data = {
        "REF": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": str(data.NAME) if data.NAME else "",
        "DURATION": str(data.DURATION) if data.DURATION else "",
        "STARTDATE": str(data.STARTDATE) if data.STARTDATE else "",
        "ENDDATE": str(data.ENDDATE) if data.ENDDATE else "",
    }
    logger.info(f"Template data for AI/ML experience letter: {template_data}")
    try:
        result = generate_document(r"C:\Users\hp\Downloads\DevRolin_ems_-System-Devrolinems\DevRolin_ems_-System-Devrolinems\Python-Backend\code\templates\Experince_AI_Template.docx", template_data, "aiml_experience_letter")
        logger.info(f"AI/ML experience letter generation result: {result}")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "AI/ML Experience letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]
        logger.info(f"AI/ML experience letter response: {response}")

        return response

    except Exception as e:
        logger.error(f"Error generating AI/ML experience letter: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate/experience-letter/webdev")
async def generate_webdev_experience_letter(data: WebDevExperienceLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Received request to generate Web Development experience letter for {data.NAME}")
    template_data = {
        "REF": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": str(data.NAME) if data.NAME else "",
        "DURATION": str(data.DURATION) if data.DURATION else "",
        "STARTDATE": str(data.STARTDATE) if data.STARTDATE else "",
        "ENDDATE": str(data.ENDDATE) if data.ENDDATE else "",
    }
    logger.info(f"Template data for Web Development experience letter: {template_data}")
    try:
        result = generate_document(r"C:\Users\hp\Downloads\DevRolin_ems_-System-Devrolinems\DevRolin_ems_-System-Devrolinems\Python-Backend\code\templates\Experience_Web_Template.docx", template_data, "webdev_experience_letter")
        logger.info(f"Web Development experience letter generation result: {result}")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Web Development Experience letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]
        logger.info(f"Web Development experience letter response: {response}")

        return response

    except Exception as e:
        logger.error(f"Error generating Web Development experience letter: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate/experience-letter/graphic-design")
async def generate_graphic_design_experience_letter(data: GraphicDesignExperienceLetterData, background_tasks: BackgroundTasks):
    logger.info(f"Received request to generate Graphic Design experience letter for {data.NAME}")
    template_data = {
        "REF": str(data.REF) if data.REF else "",
        "DATE": str(data.DATE) if data.DATE else "",
        "NAME": str(data.NAME) if data.NAME else "",
        "DURATION": str(data.DURATION) if data.DURATION else "",
        "STARTDATE": str(data.STARTDATE) if data.STARTDATE else "",
        "ENDDATE": str(data.ENDDATE) if data.ENDDATE else "",
    }
    logger.info(f"Template data for Graphic Design experience letter: {template_data}")
    try:
        result = generate_document(r"C:\Users\hp\Downloads\DevRolin_ems_-System-Devrolinems\DevRolin_ems_-System-Devrolinems\Python-Backend\code\templates\Experience_Graphic_Template.docx", template_data, "graphic_design_experience_letter")
        logger.info(f"Graphic Design experience letter generation result: {result}")
        cleanup_old_files(background_tasks)

        response = {
            "success": True,
            "message": "Graphic Design Experience letter generated successfully",
            "docx_url": f"/download/{result['docx_filename']}",
        }
        if result.get("pdf_filename"):
            response["pdf_url"] = f"/download/{result['pdf_filename']}"
        if result.get("warning"):
            response["warning"] = result["warning"]
        logger.info(f"Graphic Design experience letter response: {response}")

        return response

    except Exception as e:
        logger.error(f"Error generating Graphic Design experience letter: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{filename}")
async def download_file(filename: str):
    logger.info(f"Received download request for file: {filename}")
    file_path = os.path.join(OUTPUT_DIR, filename)
    logger.info(f"Checking file existence: {file_path}")

    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise HTTPException(status_code=404, detail="File not found")
    
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if filename.endswith(".docx") else "application/pdf"
    logger.info(f"Returning FileResponse for {filename} with content_type: {content_type}")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=content_type
    )

@app.get("/")
async def root():
    logger.info("Root endpoint accessed")
    return {"message": "Document Generator API is running. Access the API documentation at /docs"}

if __name__ == "__main__":
    logger.info("Starting FastAPI server")
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)